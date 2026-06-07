"""Word文档解析器"""

from typing import Any, Dict, List

from loguru import logger

from app.rag.parsers.base import BaseParser


class WordParser(BaseParser):
    """Word文档解析器，使用python-docx解析.docx文件"""

    def parse(self, file_path: str) -> List[Dict[str, Any]]:
        """解析Word文档，按段落/标题分割，提取表格

        Args:
            file_path: Word文件路径

        Returns:
            解析结果列表:
                - content: 段落/表格文本内容
                - metadata: 包含paragraph_index、heading_level等信息

        Raises:
            FileNotFoundError: 文件不存在
            Exception: Word解析失败
        """
        from docx import Document

        results: List[Dict[str, Any]] = []

        try:
            doc = Document(file_path)
            logger.info(f"开始解析Word文件: {file_path}，共 {len(doc.paragraphs)} 个段落")

            # 按段落解析
            current_section: List[str] = []
            current_heading: str = ""

            for para_idx, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if not text:
                    continue

                # 判断是否为标题
                heading_level = self._get_heading_level(paragraph)

                if heading_level > 0:
                    # 遇到标题，先保存之前的段落
                    if current_section:
                        section_text = self.clean_text("\n".join(current_section))
                        if section_text:
                            results.append(
                                {
                                    "content": section_text,
                                    "metadata": {
                                        "heading": current_heading,
                                        "paragraph_start": para_idx - len(current_section),
                                        "paragraph_end": para_idx - 1,
                                        "source": file_path,
                                    },
                                }
                            )
                        current_section = []

                    current_heading = text
                    results.append(
                        {
                            "content": self.clean_text(text),
                            "metadata": {
                                "heading": text,
                                "heading_level": heading_level,
                                "paragraph_index": para_idx,
                                "source": file_path,
                                "type": "heading",
                            },
                        }
                    )
                else:
                    current_section.append(text)

            # 保存最后一个段落组
            if current_section:
                section_text = self.clean_text("\n".join(current_section))
                if section_text:
                    results.append(
                        {
                            "content": section_text,
                            "metadata": {
                                "heading": current_heading,
                                "paragraph_start": len(doc.paragraphs) - len(current_section),
                                "paragraph_end": len(doc.paragraphs) - 1,
                                "source": file_path,
                            },
                        }
                    )

            # 提取表格
            table_results = self._extract_tables(doc, file_path)
            results.extend(table_results)

            logger.info(f"Word解析完成，共提取 {len(results)} 个内容块")

        except FileNotFoundError:
            raise FileNotFoundError(f"文件不存在: {file_path}")
        except Exception as e:
            logger.error(f"Word解析失败: {e}")
            raise Exception(f"Word解析失败: {e}")

        return results

    def _get_heading_level(self, paragraph) -> int:
        """获取段落的标题级别

        Args:
            paragraph: docx段落对象

        Returns:
            标题级别(1-9)，非标题返回0
        """
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading"):
            try:
                return int(style_name.replace("Heading", "").strip())
            except ValueError:
                return 1
        if style_name.startswith("标题"):
            try:
                return int(style_name.replace("标题", "").strip())
            except ValueError:
                return 1
        return 0

    def _extract_tables(self, doc, file_path: str) -> List[Dict[str, Any]]:
        """提取Word文档中的表格

        Args:
            doc: Document对象
            file_path: 文件路径

        Returns:
            表格内容列表
        """
        table_results: List[Dict[str, Any]] = []

        for table_idx, table in enumerate(doc.tables):
            try:
                rows_text: List[str] = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows_text.append(" | ".join(cells))

                table_text = "\n".join(rows_text)
                cleaned = self.clean_text(table_text)
                if cleaned:
                    table_results.append(
                        {
                            "content": cleaned,
                            "metadata": {
                                "table_index": table_idx,
                                "source": file_path,
                                "type": "table",
                            },
                        }
                    )
            except Exception as e:
                logger.warning(f"Word表格 {table_idx} 提取失败: {e}")
                continue

        return table_results
